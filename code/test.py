from docx import Document

# Create a Document object
doc = Document()
doc.add_heading('原子图构建过程（公式视角）', 0)

# Add a section for the introduction
doc.add_paragraph(
    '本文件描述了原子图构建的过程，从构建公式的角度分析每个步骤的数学表达。'
)

# Add a section for the graph construction definitions
doc.add_heading('一、原子图表示定义', level=1)
doc.add_paragraph(
    '给定一个分子 \( \\mathcal{M} \)，我们构建一个无向属性图：\n\n'
    '\\[ G = (V, E, X, E_f) \\]\n\n'
    '其中：\n'
    '- \( V = \\{v_1, v_2, \\dots, v_N\\} \\)：原子集合，表示分子的原子。\n'
    '- \( E \\subseteq V \\times V \\)：边集合，表示原子之间的化学键。\n'
    '- \( X \\in \\mathbb{Z}^{N \\times d_a} \\)：原子特征矩阵，每行是一个原子的类别型特征（如原子类型、电荷、杂环信息等）。\n'
    '- \( E_f \\in \\mathbb{Z}^{|E| \\times d_e} \\)：边特征矩阵，每行为一个键的类别型特征（如键类型、共轭性等）。'
)

# Add a section for embedding maps
doc.add_heading('二、嵌入映射', level=1)
doc.add_paragraph(
    '由于原子和键的特征是多维类别型（one-hot-like），需要映射为整数嵌入索引。\n\n'
    '1. 原子特征嵌入：\n'
    '\\[ X\' = \\text{convert\\_to\\_single\\_emb}(X) \\in \\mathbb{Z}^{N \\times d_a} \\]\n'
    '使用如下映射策略：\n'
    '\\[ X\'_i^{(k)} = X_i^{(k)} + \\sum_{j=1}^{k-1} s_j + 1 \\]\n'
    '即所有类别特征被线性展开为一个单一嵌入空间（避免索引重叠），保留 0 做 padding。\n\n'
    '2. 边特征嵌入：\n'
    '\\[ E_f\' = \\text{convert\\_to\\_single\\_emb}(E_f) \\in \\mathbb{Z}^{|E| \\times d_e} \\]'
)

# Add a section for adjacency matrix and edge feature tensor construction
doc.add_heading('三、邻接矩阵与边特征张量构造', level=1)
doc.add_paragraph(
    '定义邻接矩阵：\n'
    '\\[ A_{ij} = \\begin{cases} 1 & \\text{if } (v_i, v_j) \\in E \\ \\ 0 & \\text{otherwise} \\end{cases} \\]\n\n'
    '度向量：\n'
    '\\[ D_i = \\sum_{j} A_{ij} \\]\n\n'
    '边特征张量 \( E\' \\in \\mathbb{Z}^{N \\times N \\times d_e} \\) 构造为：\n'
    '\\[ E\'_{ij} = \\begin{cases} \\text{嵌入后的边特征} & \\text{if } A_{ij} = 1 \\ \\ 0 & \\text{otherwise} \\end{cases} \\]'
)

# Add a section for shortest path matrix
doc.add_heading('四、最短路径矩阵', level=1)
doc.add_paragraph(
    '用 Floyd-Warshall 算法在邻接图上计算最短路径矩阵：\n'
    '\\[ SP_{ij} = \\min \\left\\{ \\text{path length from } v_i \\text{ to } v_j \\right\\} \\]\n'
    '设不可达路径设置为常数 \( c = 510 \\)，则：\n'
    '\\[ SP_{ij} = \\begin{cases} \\text{最短路径长度} & \\text{if reachable} \\ \\ 510 & \\text{otherwise} \\end{cases} \\]'
)

# Add a section for Pair Type features
doc.add_heading('五、Pair Type 特征', level=1)
doc.add_paragraph(
    '设每个原子的主类特征为 \( x_i^{(0)} \\)，构造原子对特征张量：\n'
    '\\[ P_{ij} = \\left[ x_i^{(0)}, x_j^{(0)} \\right] \\]\n'
    '并用类似嵌入方法将其映射为统一索引空间。'
)

# Add a section for final feature combination
doc.add_heading('六、最终特征组合', level=1)
doc.add_paragraph(
    '如果掩码矩阵为 \( M \\in \\{0, 1\\}^N \\)，表示选择的原子子集，则最终图特征集合为：\n'
    '- 原子特征：\( \\text{atom\\_feat} = X\'[M] \\)\n'
    '- 边特征：\( \\text{edge\\_feat} = E\'[M][:, M] \\)\n'
    '- 距离矩阵：\( \\text{shortest\\_path} = SP[M][:, M] \\)\n'
    '- 原子度：\( D[M] \\)\n'
    '- 原子对类型：\( \\text{pair\\_type} = \\text{emb}(P[M][:, M]) \\)\n'
    '- 注意力偏置：初始化为零矩阵'
)

# Save the document
file_path = "原子图构建分析.docx"
doc.save(file_path)

file_path
